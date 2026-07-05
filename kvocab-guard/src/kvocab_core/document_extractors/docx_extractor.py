from __future__ import annotations

from pathlib import Path

from kvocab_core.document_extractors.base import BaseExtractor, UnsupportedFormatError
from kvocab_core.schemas import ExtractedDocument, ExtractedPage


class DocxExtractor(BaseExtractor):
    extensions = (".docx",)

    def extract(self, path: Path) -> ExtractedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise UnsupportedFormatError("python-docx is not installed") from exc
        doc = Document(str(path))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paras)
        return ExtractedDocument(text=text, pages=[ExtractedPage(page_no=1, text=text)])
